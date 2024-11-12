import sys
import os
from PyQt5.QtWidgets import (QApplication, QMainWindow, QPushButton, QLabel, 
                            QVBoxLayout, QWidget, QFileDialog, QMessageBox,
                            QListWidget, QHBoxLayout)
from PyQt5.QtGui import QPixmap
import pytesseract
from PIL import Image
from docx import Document

class ImageToWordApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Олон зураг -> Word хөрвүүлэгч')
        self.setGeometry(100, 100, 800, 600)
        
        # Tesseract тохиргоо
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        os.environ['TESSDATA_PREFIX'] = r'C:\Program Files\Tesseract-OCR\tessdata'
        
        # Үндсэн widget
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        
        # Хэвтээ layout
        self.main_layout = QHBoxLayout(self.central_widget)
        
        # Зүүн талын panel (жагсаалт)
        self.left_panel = QWidget()
        self.left_layout = QVBoxLayout(self.left_panel)
        
        # Зургийн жагсаалт
        self.image_list = QListWidget()
        self.image_list.itemSelectionChanged.connect(self.show_selected_image)
        self.left_layout.addWidget(QLabel('Сонгосон зургууд:'))
        self.left_layout.addWidget(self.image_list)
        
        # Товчнууд
        self.select_button = QPushButton('Зураг нэмэх')
        self.select_button.clicked.connect(self.select_images)
        self.left_layout.addWidget(self.select_button)
        
        self.remove_button = QPushButton('Зураг хасах')
        self.remove_button.clicked.connect(self.remove_image)
        self.left_layout.addWidget(self.remove_button)
        
        self.convert_button = QPushButton('Word болгох')
        self.convert_button.clicked.connect(self.convert_to_word)
        self.left_layout.addWidget(self.convert_button)
        
        # Баруун талын panel (урьдчилан харах)
        self.right_panel = QWidget()
        self.right_layout = QVBoxLayout(self.right_panel)
        
        self.preview_label = QLabel('Зураг сонгоогүй байна')
        self.preview_label.setMinimumSize(400, 400)
        self.right_layout.addWidget(QLabel('Урьдчилан харах:'))
        self.right_layout.addWidget(self.preview_label)
        
        # Панелуудыг нэмэх
        self.main_layout.addWidget(self.left_panel)
        self.main_layout.addWidget(self.right_panel)
        
        # Зургийн замуудыг хадгалах жагсаалт
        self.image_paths = []
        
    def select_images(self):
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "Зураг сонгох",
            "",
            "Зургийн файлууд (*.png *.jpg *.jpeg)"
        )
        
        if files:
            for file_path in files:
                if file_path not in self.image_paths:
                    self.image_paths.append(file_path)
                    self.image_list.addItem(os.path.basename(file_path))
    
    def remove_image(self):
        current_row = self.image_list.currentRow()
        if current_row >= 0:
            self.image_list.takeItem(current_row)
            self.image_paths.pop(current_row)
            if not self.image_paths:
                self.preview_label.setText('Зураг сонгоогүй байна')
    
    def show_selected_image(self):
        current_row = self.image_list.currentRow()
        if current_row >= 0:
            pixmap = QPixmap(self.image_paths[current_row])
            scaled_pixmap = pixmap.scaled(400, 400)
            self.preview_label.setPixmap(scaled_pixmap)
    
    def convert_to_word(self):
        if not self.image_paths:
            QMessageBox.warning(self, 'Анхааруулга', 'Эхлээд зураг сонгоно уу!')
            return
            
        try:
            # Word файл үүсгэх
            doc = Document()
            
            # Зураг бүрийг боловсруулах
            for image_path in self.image_paths:
                img = Image.open(image_path)
                text = pytesseract.image_to_string(img, lang='mon')
                
                # Файлын нэрийг гарчиг болгон оруулах
                doc.add_heading(os.path.basename(image_path), level=1)
                doc.add_paragraph(text)
                doc.add_paragraph('\n')  # Хоосон зай нэмэх
            
            # Хадгалах
            save_path, _ = QFileDialog.getSaveFileName(
                self,
                "Word файл хадгалах",
                "output.docx",
                "Word файлууд (*.docx)"
            )
            
            if save_path:
                doc.save(save_path)
                QMessageBox.information(
                    self, 
                    'Амжилттай', 
                    f'Бүх текстийг {save_path} файлд амжилттай хадгаллаа.'
                )
                
        except pytesseract.TesseractNotFoundError:
            QMessageBox.critical(
                self, 
                'Алдаа', 
                'Tesseract OCR олдсонгүй. Зам зөв эсэхийг шалгана уу'
            )
        except pytesseract.TesseractError as e:
            QMessageBox.critical(
                self, 
                'Алдаа',
                f'Tesseract алдаа: {str(e)}\nМонгол хэлний файлууд суусан эсэхийг шалгана уу'
            )
        except Exception as e:
            QMessageBox.critical(self, 'Алдаа', f'Тодорхойгүй алдаа: {str(e)}')

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = ImageToWordApp()
    window.show()
    sys.exit(app.exec_()) 